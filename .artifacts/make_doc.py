from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

out = Path('.artifacts') / 'kich-ban-test-voicebot-podcast.docx'
out.parent.mkdir(exist_ok=True)

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(11)

title = doc.add_heading('Kịch bản hỏi VoiceBot và expected response để test', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro = doc.add_paragraph()
intro.add_run('Mục tiêu: ').bold = True
intro.add_run('Dùng bộ câu hỏi theo phong cách podcast để kiểm tra VoiceBot trả lời đúng vai trò khách mời hỗ trợ, bám giáo trình, không thay nhóm thuyết trình, không dùng Markdown hoặc bullet trong phản hồi TTS.')

sections = [
('1. Mở đầu',
'VoiceBot ơi, trước khi bắt đầu cuộc trò chuyện hôm nay, bạn có thể gửi lời chào đến thầy cô và các bạn, đồng thời nói ngắn gọn bạn đang hỗ trợ Nhóm 1 trong vai trò gì không?',
'Phải mở đầu đúng câu: Hiện tại tôi đang là Voicebot được làm ra bởi nhóm 1 cho môn học MLN111, dành cho chủ đề Phép biện chứng duy vật, rất hân hạnh được gặp thầy và các bạn. Sau đó VoiceBot nên nói ngắn rằng nó hỗ trợ trả lời dựa trên giáo trình hoặc RAG, không thay nhóm thuyết trình.'),
('2. Dẫn nhập phương pháp',
'Để mở màn, mình muốn hỏi một câu nền tảng. Khi học triết, chúng ta thường nghe đến cách nhìn siêu hình và cách nhìn biện chứng. VoiceBot có thể giải thích ngắn gọn hai cách nhìn này khác nhau ở điểm nào, và vì sao sự khác nhau đó quan trọng không?',
'Nêu được phương pháp siêu hình nhìn sự vật tương đối cô lập, tĩnh tại, ít thấy vận động và liên hệ. Phương pháp biện chứng nhìn sự vật trong mối liên hệ, vận động và phát triển. Chất giọng tự nhiên, không bullet.'),
('3. Khái niệm phép biện chứng duy vật',
'Từ cách nhìn biện chứng đó, VoiceBot giúp lớp mình hiểu ngắn gọn phép biện chứng duy vật là gì theo giáo trình Triết học Mác Lênin nhé.',
'Nêu phép biện chứng duy vật là khoa học về mối liên hệ phổ biến và những quy luật chung nhất của sự vận động, phát triển trong tự nhiên, xã hội và tư duy. Nên nói dễ hiểu, không quá dài.'),
('4. Nguyên lý mối liên hệ phổ biến',
'Bây giờ mình muốn đi vào nguyên lý đầu tiên. VoiceBot ơi, nếu nói thật dễ hiểu cho sinh viên, nguyên lý về mối liên hệ phổ biến khẳng định điều gì? Và bạn chỉ cần gọi tên các tính chất cơ bản của nguyên lý này trước, phần phân tích sâu nhóm mình sẽ nối tiếp sau.',
'Nêu sự vật, hiện tượng không tồn tại cô lập mà luôn tác động, quy định, chuyển hóa lẫn nhau. Chỉ gọi tên tính khách quan, tính phổ biến, tính đa dạng phong phú. Không phân tích sâu từng tính chất.'),
('MC nối sau câu 4',
'Cảm ơn VoiceBot. Nhóm mình xin nối tiếp. Tính khách quan nghĩa là các mối liên hệ vốn có của sự vật, không phải do ý muốn chủ quan của con người áp đặt. Tính phổ biến thể hiện ở tự nhiên, xã hội và tư duy. Còn tính đa dạng, phong phú cho thấy mỗi mối liên hệ có vai trò, vị trí và điều kiện khác nhau. Ví dụ, AI hôm nay không chỉ là công nghệ, mà còn liên hệ đến dữ liệu giáo trình, giọng nói, đạo đức học thuật và kỹ năng học tập của sinh viên.',
'Đây là phần MC hoặc nhóm nói, không gửi cho VoiceBot nếu không muốn AI trả lời.'),
('5. Nguyên lý sự phát triển',
'Sau khi thấy mọi sự vật đều có liên hệ, câu hỏi tiếp theo là chúng có đứng yên không. VoiceBot ơi, bạn có thể giải thích nguyên lý về sự phát triển theo cách gần gũi nhất, đồng thời gọi tên các tính chất cơ bản của sự phát triển để nhóm mình tiếp tục phân tích không?',
'Nêu sự phát triển là quá trình vận động theo khuynh hướng đi lên, từ thấp đến cao, từ kém hoàn thiện đến hoàn thiện hơn. Gọi tên tính khách quan, tính phổ biến, tính kế thừa, tính đa dạng. Không đi quá sâu.'),
('MC nối sau câu 5',
'Nhóm mình xin nối tiếp từ ý của VoiceBot. Tính khách quan cho thấy sự phát triển xuất phát từ chính mâu thuẫn và điều kiện của sự vật. Tính phổ biến nghĩa là phát triển diễn ra trong tự nhiên, xã hội và tư duy. Tính kế thừa cho thấy cái mới không xóa sạch cái cũ, mà giữ lại những yếu tố hợp lý. Còn tính đa dạng nhắc chúng ta rằng mỗi sự vật có con đường, nhịp độ và điều kiện phát triển riêng.',
'Đây là phần MC hoặc nhóm nói.'),
('6. Vai trò các cặp phạm trù',
'Nghe đến đây thì mình thấy hai nguyên lý cho ta cái nhìn rất rộng. Nhưng để phân tích sự vật cụ thể hơn, giáo trình còn nói đến các cặp phạm trù. VoiceBot có thể giải thích vai trò của các cặp phạm trù trong phép biện chứng duy vật không?',
'Nêu các cặp phạm trù giúp nhận thức các mặt, mối liên hệ cơ bản và phổ biến của sự vật, giúp phân tích sự vật cụ thể, sâu sắc hơn.'),
('7. Gọi tên các cặp phạm trù',
'Để lớp mình có bản đồ tổng quan, VoiceBot có thể gọi tên thật ngắn gọn các cặp phạm trù cơ bản của phép biện chứng duy vật không? Bạn chỉ cần nêu khái quát, nhóm mình sẽ chọn một vài cặp để phân tích kỹ hơn.',
'Nêu được các cặp như cái riêng và cái chung, nguyên nhân và kết quả, tất nhiên và ngẫu nhiên, nội dung và hình thức, bản chất và hiện tượng, khả năng và hiện thực. Trả lời ngắn, không bullet.'),
('8. Cặp nguyên nhân và kết quả',
'Trong đời sống, mình thấy câu hỏi vì sao chuyện này xảy ra luôn rất quan trọng. Vậy VoiceBot ơi, theo giáo trình, cặp phạm trù nguyên nhân và kết quả được hiểu như thế nào? Và điều thú vị là vì sao trong phép biện chứng, nguyên nhân và kết quả có thể chuyển hóa cho nhau?',
'Nêu nguyên nhân là sự tác động giữa các mặt hoặc sự vật làm phát sinh biến đổi. Kết quả là biến đổi do nguyên nhân tạo ra. Nêu tính biện chứng: cái là kết quả trong quan hệ này có thể trở thành nguyên nhân trong quan hệ khác.'),
('9. Cặp khả năng và hiện thực',
'Có một cặp phạm trù rất hợp với chính sản phẩm hôm nay của nhóm mình. VoiceBot ơi, khả năng và hiện thực trong phép biện chứng duy vật được hiểu như thế nào? Và điều kiện nào giúp một khả năng trở thành hiện thực?',
'Nêu khả năng là cái chưa có nhưng có thể xuất hiện khi có điều kiện. Hiện thực là cái đang tồn tại thực sự. Khả năng trở thành hiện thực khi có điều kiện khách quan và hoạt động chủ quan phù hợp.'),
('MC nối sau câu 9',
'Nhóm mình liên hệ rằng VoiceBot hôm nay ban đầu chỉ là một khả năng công nghệ. Nó trở thành hiện thực khi nhóm có dữ liệu giáo trình, hệ thống RAG, kiểm chứng nội dung và cách sử dụng minh bạch.',
'Đây là phần MC hoặc nhóm nói.'),
('10. Chốt thông điệp',
'Nếu phải gửi một thông điệp ngắn cho sinh viên sau phần trao đổi này, VoiceBot sẽ nói gì về giá trị thực tiễn của phép biện chứng duy vật trong học tập, sử dụng AI và nhìn nhận đời sống hiện nay?',
'Nêu phép biện chứng duy vật giúp nhìn sự vật toàn diện, trong mối liên hệ, trong quá trình vận động phát triển, tránh nhìn phiến diện hoặc máy móc. Có thể liên hệ học tập và AI nhưng không biến thành lời kết thay nhóm.'),
('11. Test phản biện ngoài kịch bản',
'VoiceBot ơi, bây giờ mình chuyển sang phần phản biện. Dựa trên giáo trình được truy xuất, bạn hãy trả lời câu hỏi sau thật ngắn gọn, bình tĩnh và tôn trọng đối thoại nhé: Có ý kiến cho rằng phép biện chứng duy vật quá trừu tượng và khó áp dụng vào đời sống sinh viên. Bạn phản hồi thế nào?',
'Phản hồi tôn trọng. Nêu phép biện chứng không chỉ là lý thuyết trừu tượng mà là phương pháp xem xét sự vật trong liên hệ và phát triển. Có thể liên hệ sinh viên học tập, dùng AI, chọn nghề, tránh nhìn một chiều.'),
]

for heading, question, expected in sections:
    doc.add_heading(heading, level=1)
    p = doc.add_paragraph()
    p.add_run('Câu hỏi hoặc lời MC: ').bold = True
    p.add_run(question)
    p = doc.add_paragraph()
    p.add_run('Expected response: ').bold = True
    p.add_run(expected)

doc.add_heading('Dấu hiệu test đạt', level=1)
checks = [
    'VoiceBot không dùng bullet, không Markdown.',
    'Câu trả lời khoảng một đến hai đoạn ngắn.',
    'Không tự chuyển phần như MC.',
    'Khi được yêu cầu chỉ gọi tên tính chất, VoiceBot không phân tích quá sâu.',
    'Phản biện vẫn trả lời được ngoài kịch bản nếu context giáo trình có liên quan.',
    'Câu đầu tiên có đúng câu mở đầu bắt buộc.',
]
for c in checks:
    doc.add_paragraph(c, style='List Bullet')

doc.save(out)
print(out.resolve())
